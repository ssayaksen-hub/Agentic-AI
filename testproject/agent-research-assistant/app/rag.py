from pathlib import Path

import chromadb
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.config import VECTOR_COLLECTION_NAME

CHROMA_PATH = "./chroma_db"


def _get_embeddings() -> OpenAIEmbeddings:
    """Create embeddings client from environment (expects OPENAI_API_KEY)."""
    return OpenAIEmbeddings()


def _load_txt(file_path: str) -> list[Document]:
    text = Path(file_path).read_bytes().decode("utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_docx(file_path: str) -> list[Document]:
    import docx  # python-docx

    doc = docx.Document(file_path)
    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_text_fallback(file_path: str) -> list[Document]:
    """Best-effort loader for any text-like file type."""
    raw = Path(file_path).read_bytes()
    if not raw:
        return []

    if b"\x00" in raw[:4096]:
        return []

    text = raw.decode("utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": file_path})]


def _extract_documents_with_fallback(file_path: str) -> list[Document]:
    """Load PDF text with PyPDFLoader, then fallback to pypdf if needed."""
    loader = PyPDFLoader(file_path)
    documents = loader.load()

    has_text = any((doc.page_content or "").strip() for doc in documents)
    if has_text:
        return documents

    reader = PdfReader(file_path)
    fallback_docs: list[Document] = []
    for page_idx, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        fallback_docs.append(
            Document(
                page_content=text,
                metadata={"source": file_path, "page": page_idx},
            )
        )

    return fallback_docs


def _load_any_document(file_path: str) -> list[Document]:
    """Load a document file into LangChain Documents."""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return _extract_documents_with_fallback(file_path)
    if suffix == ".txt":
        return _load_txt(file_path)
    if suffix == ".docx":
        return _load_docx(file_path)
    return _load_text_fallback(file_path)


def add_document_to_vectorstore(
    file_path: str,
    collection_name: str = VECTOR_COLLECTION_NAME,
):
    """Load a document, split it into chunks, and append it to Chroma."""
    doc_path = Path(file_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    documents = _load_any_document(file_path)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
    )
    docs = text_splitter.split_documents(documents)

    docs = [doc for doc in docs if doc.page_content and doc.page_content.strip()]
    if not docs:
        raise ValueError(
            "No non-empty text extracted from the document. "
            "Ensure the file contains readable text content."
        )

    vectorstore = Chroma(
        persist_directory=CHROMA_PATH,
        embedding_function=_get_embeddings(),
        collection_name=collection_name,
    )

    vectorstore.add_documents(docs)
    if hasattr(vectorstore, "persist"):
        vectorstore.persist()

    return vectorstore


def create_vectorstore(
    file_path: str = "data/sample.pdf",
    collection_name: str = VECTOR_COLLECTION_NAME,
):
    """Backwards-compatible wrapper around add_document_to_vectorstore."""
    return add_document_to_vectorstore(file_path, collection_name)


def load_vectorstore(collection_name: str = VECTOR_COLLECTION_NAME):
    """Load an existing Chroma vector store from disk."""
    return Chroma(
        persist_directory=CHROMA_PATH,
        embedding_function=_get_embeddings(),
        collection_name=collection_name,
    )


def rebuild_vectorstore(
    data_dir: str = "data",
    collection_name: str = VECTOR_COLLECTION_NAME,
):
    """Rebuild the vector store from all files in the data directory."""
    data_path = Path(data_dir)
    doc_files = sorted(
        f for f in data_path.iterdir()
        if f.is_file() and not f.name.startswith(".")
    )

    client = chromadb.PersistentClient(path=CHROMA_PATH)
    try:
        client.delete_collection(collection_name)
    except Exception:  # noqa: BLE001
        pass

    if not doc_files:
        return None

    documents: list[Document] = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
    )

    for doc in doc_files:
        try:
            loaded_docs = _load_any_document(str(doc))
        except Exception:  # noqa: BLE001
            continue
        split_docs = text_splitter.split_documents(loaded_docs)
        documents.extend(
            chunk for chunk in split_docs if chunk.page_content and chunk.page_content.strip()
        )

    if not documents:
        raise ValueError(
            "No non-empty text extracted from files in data/. "
            "Ensure the files contain readable text content."
        )

    vectorstore = Chroma.from_documents(
        documents,
        _get_embeddings(),
        collection_name=collection_name,
        persist_directory=CHROMA_PATH,
    )
    if hasattr(vectorstore, "persist"):
        vectorstore.persist()

    return vectorstore
